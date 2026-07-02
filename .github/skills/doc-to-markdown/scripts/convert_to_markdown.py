#!/usr/bin/env python3
"""
convert_to_markdown.py
Converts DOCX, PDF, and VTT files to Markdown (.md).

Usage:
    python convert_to_markdown.py <input_path> [--output-dir <dir>] [--vtt-speaker-labels] [--delete-source]

Arguments:
    input_path          Path to a single file or a directory of files to convert.
    --output-dir        Directory to write output .md files (defaults to same dir as input).
    --vtt-speaker-labels  When set, preserves speaker labels from VTT in the output.
    --delete-source     Delete the original file after successful conversion.

Dependencies:
    pip install python-docx pdfplumber
"""

import argparse
import re
import sys
from pathlib import Path


def convert_docx(input_path: Path, output_path: Path) -> None:
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx is not installed. Run: pip install python-docx")
        sys.exit(1)

    document = Document(str(input_path))
    lines = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            lines.append("")
            continue

        style = paragraph.style.name if paragraph.style else ""

        if style.startswith("Heading 1"):
            lines.append(f"# {text}")
        elif style.startswith("Heading 2"):
            lines.append(f"## {text}")
        elif style.startswith("Heading 3"):
            lines.append(f"### {text}")
        elif style.startswith("Heading 4"):
            lines.append(f"#### {text}")
        elif style.startswith("List"):
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for table in document.tables:
        header_row = table.rows[0].cells if table.rows else []
        if header_row:
            headers = [cell.text.strip() for cell in header_row]
            lines.append("")
            lines.append("| " + " | ".join(headers) + " |")
            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
            lines.append("")

    output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
    print(f"Converted DOCX -> {output_path}")


def parse_vtt_timestamp(timestamp: str) -> str:
    """Return a simplified HH:MM:SS timestamp string."""
    return timestamp.split(".")[0]


def convert_vtt(input_path: Path, output_path: Path, keep_speaker_labels: bool = True) -> None:
    raw = input_path.read_text(encoding="utf-8", errors="replace")
    lines_out = []
    current_speaker = None
    current_text_lines = []

    # VTT blocks are separated by blank lines
    blocks = re.split(r"\n{2,}", raw.strip())

    for block in blocks:
        block_lines = block.strip().splitlines()
        if not block_lines:
            continue
        # Skip the WEBVTT header line
        if block_lines[0].strip().upper().startswith("WEBVTT"):
            continue
        # Skip cue identifier lines (pure digits or alphanumeric IDs with no arrows)
        if len(block_lines) == 1 and re.match(r"^[\w\-]+$", block_lines[0]):
            continue

        # Find timestamp line
        timestamp_line = None
        text_lines = []
        for line in block_lines:
            if "-->" in line:
                timestamp_line = line
            elif timestamp_line is not None:
                text_lines.append(line)

        if not timestamp_line or not text_lines:
            continue

        start_ts = parse_vtt_timestamp(timestamp_line.split("-->")[0].strip())
        full_text = " ".join(text_lines).strip()

        # Detect and strip speaker labels like "Speaker Name: " or "<v Speaker>..."
        speaker = None
        tag_match = re.match(r"<v\s+([^>]+)>(.*)", full_text, re.DOTALL)
        colon_match = re.match(r"^([A-Z][^:]{0,40}):\s+(.*)", full_text)

        if tag_match:
            speaker = tag_match.group(1).strip()
            full_text = tag_match.group(2).strip()
        elif colon_match:
            speaker = colon_match.group(1).strip()
            full_text = colon_match.group(2).strip()

        # Clean HTML tags remaining in text
        full_text = re.sub(r"<[^>]+>", "", full_text).strip()

        if not full_text:
            continue

        if keep_speaker_labels and speaker:
            if speaker != current_speaker:
                if current_text_lines:
                    lines_out.append(" ".join(current_text_lines))
                    lines_out.append("")
                lines_out.append(f"**{speaker}** *{start_ts}*")
                current_speaker = speaker
                current_text_lines = [full_text]
            else:
                current_text_lines.append(full_text)
        else:
            lines_out.append(f"*{start_ts}* {full_text}")

    if current_text_lines:
        lines_out.append(" ".join(current_text_lines))

    output_path.write_text("\n".join(lines_out).strip() + "\n", encoding="utf-8")
    print(f"Converted VTT  -> {output_path}")


def convert_pdf(input_path: Path, output_path: Path) -> bool:
    """
    Convert PDF to Markdown, extracting text, tables, and images.
    Returns True if images were extracted, False otherwise.
    """
    try:
        import pdfplumber
    except ImportError:
        print("ERROR: pdfplumber is not installed. Run: pip install pdfplumber")
        sys.exit(1)

    lines = []
    images_extracted = False
    image_dir = output_path.parent / (output_path.stem + "_images")
    image_counter = 0

    try:
        with pdfplumber.open(str(input_path)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text with layout awareness
                text = page.extract_text()
                if text:
                    # Split into paragraphs and detect potential headings
                    paragraphs = text.split("\n")
                    for para in paragraphs:
                        para = para.strip()
                        if not para:
                            lines.append("")
                            continue
                        
                        # Simple heuristic: if text is very short and appears at top of page,
                        # or in first paragraph, it's likely a heading
                        # For now, we'll mark short lines (< 60 chars) at the start as potential headings
                        if len(para) < 60 and page_num == 1 and len(lines) < 5:
                            lines.append(f"# {para}")
                        else:
                            lines.append(para)
                
                # Extract tables
                tables = page.extract_tables()
                if tables:
                    lines.append("")
                    for table in tables:
                        if table:
                            # Convert table to pipe format
                            header_row = table[0]
                            headers = [str(cell) if cell else "" for cell in header_row]
                            lines.append("| " + " | ".join(headers) + " |")
                            lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                            
                            for row in table[1:]:
                                cells = [str(cell) if cell else "" for cell in row]
                                lines.append("| " + " | ".join(cells) + " |")
                            lines.append("")
                
                # Extract images
                for img in page.images:
                    try:
                        image_dir.mkdir(parents=True, exist_ok=True)
                        bbox = img["top"], img["bottom"], img["x0"], img["x1"]
                        im = page.crop(bbox).to_image()
                        image_path = image_dir / f"image_{image_counter}.png"
                        im.save(image_path)
                        
                        # Add image reference to markdown
                        lines.append(f"![Image](image_{image_counter}.png)")
                        image_counter += 1
                        images_extracted = True
                    except Exception as e:
                        print(f"Warning: Could not extract image from page {page_num}: {e}")
    
    except Exception as e:
        print(f"ERROR: Failed to process PDF {input_path}: {e}")
        sys.exit(1)

    if lines:
        output_path.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")
        print(f"Converted PDF  -> {output_path}")
        if images_extracted:
            print(f"  Images extracted to: {image_dir}")
    else:
        print(f"WARNING: No content extracted from {input_path}")

    return images_extracted


def main():
    parser = argparse.ArgumentParser(description="Convert PDF/DOCX/VTT files to Markdown.")
    parser.add_argument("input_path", help="File or directory to convert.")
    parser.add_argument("--output-dir", default=None, help="Output directory for .md files.")
    parser.add_argument("--vtt-speaker-labels", action="store_true", default=True,
                        help="Preserve speaker labels in VTT output (default: on).")
    parser.add_argument("--delete-source", action="store_true", default=False,
                        help="Delete the original file after successful conversion.")
    args = parser.parse_args()

    input_path = Path(args.input_path)
    output_dir = Path(args.output_dir) if args.output_dir else None

    if input_path.is_dir():
        files = list(input_path.glob("*.docx")) + list(input_path.glob("*.vtt")) + list(input_path.glob("*.pdf"))
    elif input_path.is_file():
        files = [input_path]
    else:
        print(f"ERROR: {input_path} is not a valid file or directory.")
        sys.exit(1)

    if not files:
        print("No .docx, .vtt, or .pdf files found.")
        sys.exit(0)

    for file in files:
        dest_dir = output_dir if output_dir else file.parent
        dest_dir.mkdir(parents=True, exist_ok=True)
        output_file = dest_dir / (file.stem + ".md")

        images_extracted = False
        
        if file.suffix.lower() == ".docx":
            convert_docx(file, output_file)
        elif file.suffix.lower() == ".vtt":
            convert_vtt(file, output_file, keep_speaker_labels=args.vtt_speaker_labels)
        elif file.suffix.lower() == ".pdf":
            images_extracted = convert_pdf(file, output_file)
        
        # Delete source file if --delete-source flag is set
        if args.delete_source:
            if images_extracted:
                print(f"  Note: Images extracted to {output_file.parent / (output_file.stem + '_images')}; "
                      f"deleting {file.name}")
            try:
                file.unlink()
                print(f"Deleted source file: {file}")
            except Exception as e:
                print(f"WARNING: Could not delete {file}: {e}")

    print(f"\nDone. {len(files)} file(s) converted.")


if __name__ == "__main__":
    main()
